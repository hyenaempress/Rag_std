import os
import re
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import docx

try:
    from pykospacing import spacing
    PYKOSPACING_AVAILABLE = True
except ImportError:
    PYKOSPACING_AVAILABLE = False
    print("PyKoSpacing 라이브러리가 설치되지 않음. 기본 띄어쓰기 복원 방법을 사용합니다.")

try:
    from konlpy.tag import Okt
    KONLPY_AVAILABLE = True
except ImportError:
    KONLPY_AVAILABLE = False

class SimpleRAGEngine:
    def __init__(self):
        self.documents = []  # 메모리에 문서 저장
        print("간단한 키워드 기반 RAG 엔진을 사용합니다.")
    
    def add_text_document(self, text, title="문서"):
        """텍스트를 청크로 나누어 저장"""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  # 더 작은 청크로 세밀하게 분할
            chunk_overlap=200,  # 오버랩 증가로 내용 누락 방지
            separators=["\n\n", "\n", ". ", "! ", "? ", "。", "，", " "]
        )
        
        doc = Document(page_content=text, metadata={"source": title})
        chunks = text_splitter.split_documents([doc])
        self.documents.extend(chunks)
        
        return len(chunks)
    
    def add_file_document(self, file_path):
        """파일을 로드하여 문서로 추가"""
        try:
            if file_path.endswith('.pdf'):
                loader = PyPDFLoader(file_path)
                documents = loader.load()
            elif file_path.endswith('.txt'):
                loader = TextLoader(file_path, encoding='utf-8')
                documents = loader.load()
            elif file_path.endswith('.docx'):
                doc = docx.Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                documents = [Document(page_content=text, metadata={"source": file_path})]
            else:
                return False, "지원하지 않는 파일 형식입니다."
            
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,  # 더 작은 청크로 세밀하게 분할
                chunk_overlap=200,  # 오버랩 증가로 내용 누락 방지
                separators=["\n\n", "\n", ". ", "! ", "? ", "。", "，", " "]
            )
            chunks = text_splitter.split_documents(documents)
            self.documents.extend(chunks)
            
            return True, f"{len(chunks)}개 청크가 추가되었습니다."
            
        except Exception as e:
            return False, f"파일 처리 중 오류: {str(e)}"
    
    def search_documents(self, query, k=3):
        """키워드 기반 문서 검색 (개선된 버전)"""
        if not self.documents:
            return []
        
        query_words = self._extract_keywords(query.lower())
        print(f"[DEBUG] 검색 키워드: {query_words}")
        scored_docs = []
        
        for doc in self.documents:
            content = doc.page_content.lower()
            score = self._calculate_relevance_score(content, query_words, query.lower())
            
            if score > 0:
                scored_docs.append((doc, score))
        
        # 점수순 정렬 후 상위 k개 반환
        scored_docs.sort(key=lambda x: x[1], reverse=True)
        
        # 디버그: 상위 3개 결과 출력
        if scored_docs:
            print(f"[DEBUG] 상위 3개 검색 결과:")
            for i, (doc, score) in enumerate(scored_docs[:3]):
                preview = doc.page_content[:100].replace('\n', ' ')
                print(f"  {i+1}. 점수: {score}, 내용: {preview}...")
        
        return [doc for doc, score in scored_docs[:k]]
    
    def _extract_keywords(self, query):
        """쿼리에서 의미있는 키워드 추출 (개선된 버전)"""
        # 불용어 제거 (질문 단어는 제외)
        stop_words = {
            '이', '그', '저', '의', '가', '을', '를', '에', '와', '과', '으로', '로',
            '은', '는', '이다', '다', '하다', '되다', '있다', '없다', '같다',
            '지', '까', '니', '냐', '야', '어', '아', '요', '에요',
            'what', 'is', 'are', 'the', 'a', 'an', 'and', 'or', 'but'
        }
        
        # 질문 단어들 (의미가 있는 경우가 많음)
        question_words = {'뭐', '뭔', '뭐야', '무엇', '어떤', '어디', '언제', '왜', '어떻게'}
        
        words = query.split()
        meaningful_words = []
        
        for word in words:
            # 정리된 단어 생성
            clean_word = word.rstrip('지?!.,').lower()
            
            # 3글자 이상이면 항상 포함 (핵심 키워드일 가능성 높음)
            if len(clean_word) >= 3:
                meaningful_words.append(clean_word)
            # 2글자이고 불용어가 아닌 경우
            elif len(clean_word) == 2 and clean_word not in stop_words:
                meaningful_words.append(clean_word)
            # 질문 단어인 경우 (컨텍스트에 따라 의미가 있을 수 있음)
            elif clean_word in question_words and len(meaningful_words) == 0:
                meaningful_words.append(clean_word)
        
        # 최소 하나의 키워드는 있어야 함
        if not meaningful_words and words:
            # 가장 긴 단어를 선택
            longest_word = max(words, key=len).rstrip('지?!.,').lower()
            if len(longest_word) > 0:
                meaningful_words.append(longest_word)
        
        return meaningful_words
    
    def _calculate_relevance_score(self, content, keywords, full_query):
        """관련성 점수 계산 (개선된 알고리즘)"""
        score = 0
        
        # 1. 완전한 구문 일치 (가장 높은 점수)
        if full_query in content:
            score += len(full_query) * 5
        
        # 2. 개별 키워드 점수 (키워드 중요도 가중치 적용)
        for keyword in keywords:
            if keyword in content:
                # 키워드 빈도
                frequency = content.count(keyword)
                
                # 키워드 길이에 따른 가중치 (긴 키워드가 더 중요)
                length_weight = len(keyword)
                
                # 핵심 키워드 보너스 (5글자 이상은 매우 중요한 키워드로 간주)
                if len(keyword) >= 5:
                    keyword_score = frequency * length_weight * 10  # 높은 가중치
                elif len(keyword) >= 3:
                    keyword_score = frequency * length_weight * 3
                else:
                    keyword_score = frequency * length_weight * 1
                    
                score += keyword_score
                
                # 키워드가 문장 시작 부분에 있으면 추가 점수
                if content[:50].count(keyword) > 0:
                    score += length_weight * 5
        
        # 3. 키워드 근접성 보너스 (키워드들이 가까이 있으면 추가 점수)
        if len(keywords) > 1:
            proximity_bonus = self._calculate_proximity_bonus(content, keywords)
            score += proximity_bonus
        
        return score
    
    def _calculate_proximity_bonus(self, content, keywords):
        """키워드 간 근접성 계산"""
        try:
            positions = {}
            for keyword in keywords:
                positions[keyword] = [m.start() for m in re.finditer(keyword, content)]
            
            # 모든 키워드가 100자 이내에 있으면 보너스
            for pos_list in positions.values():
                for pos in pos_list:
                    nearby_keywords = 0
                    for other_keyword, other_positions in positions.items():
                        for other_pos in other_positions:
                            if abs(pos - other_pos) <= 100:  # 100자 이내
                                nearby_keywords += 1
                    
                    if nearby_keywords >= len(keywords):
                        return 50  # 근접성 보너스
            
            return 0
        except:
            return 0
    
    def get_rag_response(self, query):
        """RAG 기반 응답 생성 (안전한 버전)"""
        try:
            relevant_docs = self.search_documents(query, k=3)
            
            if not relevant_docs:
                return "죄송합니다. 관련된 문서를 찾을 수 없습니다."
            
            # 간단하고 안전한 응답 생성
            return self._generate_safe_response(query, relevant_docs)
            
        except Exception as e:
            print(f"RAG 응답 생성 오류: {e}")
            return f"검색은 완료했지만 응답 처리 중 문제가 발생했습니다. 오류: {str(e)}"
    
    def _generate_safe_response(self, query, docs):
        """안전한 응답 생성 (개선된 버전)"""
        try:
            response_parts = [f'💡 **"{query}"**에 대한 답변:\n']
            
            # 최대 2개의 관련 문서 사용
            num_docs_to_use = min(2, len(docs))
            combined_content = []
            sources = []
            
            for i in range(num_docs_to_use):
                doc = docs[i]
                content = doc.page_content
                source = doc.metadata.get('source', '문서')
                
                # 텍스트 정리
                clean_content = self._simple_clean_text(content)
                
                # 각 문서에서 500자씩 가져오기
                if len(clean_content) > 500:
                    clean_content = clean_content[:500]
                
                combined_content.append(clean_content)
                if source not in sources:
                    sources.append(source)
            
            # 출처 표시
            if sources:
                source_names = [s.split('\\')[-1] if '\\' in s else s for s in sources]
                response_parts.append(f"**📚 출처:** {', '.join(source_names[:2])}")
            
            # 내용 결합 (최대 800자)
            full_content = '\n\n'.join(combined_content)
            if len(full_content) > 800:
                full_content = full_content[:800] + "..."
            
            response_parts.append(f"**내용:**\n{full_content}")
            
            # 추가 안내
            if len(docs) > 2:
                response_parts.append(f"\n📄 추가로 {len(docs)-2}개의 관련 문서가 더 있습니다.")
            
            response_parts.append("\n💡 더 구체적인 질문을 해주시면 더 정확한 답변을 드릴게요!")
            
            return "\n".join(response_parts)
            
        except Exception as e:
            print(f"안전한 응답 생성 오류: {e}")
            # 최후의 수단 - 아주 간단한 응답
            try:
                content = docs[0].page_content[:200] + "..."
                return f'💡 **"{query}"**에 대한 답변:\n\n{content}\n\n💡 더 자세한 내용이 필요하시면 다시 질문해주세요!'
            except:
                return "문서를 찾았지만 응답을 생성하는 중 오류가 발생했습니다."
    
    def _simple_clean_text(self, text):
        """간단한 텍스트 정리 (띄어쓰기 복원 포함)"""
        try:
            # 기본 정리
            text = re.sub(r'\s+', ' ', text)  # 연속 공백 제거
            text = text.strip()
            
            # 고급 띄어쓰기 복원 함수 사용
            text = self.restore_korean_spacing(text)
            
            return text
        except Exception as e:
            print(f"텍스트 정리 중 오류: {e}")
            # 오류 발생시 기본 띄어쓰기만 적용
            text = re.sub(r'([.!?])([가-힣A-Za-z])', r'\1 \2', text)
            text = re.sub(r'([A-Za-z])([가-힣])', r'\1 \2', text)
            text = re.sub(r'([가-힣])([A-Za-z])', r'\1 \2', text)
            return text
    
    def _extract_key_summary(self, content, query):
        """핵심 요약 추출"""
        try:
            # 쿼리 키워드가 포함된 문장들 찾기
            sentences = self._split_into_sentences(content)
            query_words = self._extract_keywords(query.lower())
            
            relevant_sentences = []
            for sentence in sentences:
                sentence_lower = sentence.lower()
                if any(word in sentence_lower for word in query_words):
                    relevant_sentences.append(sentence.strip())
                    if len(relevant_sentences) >= 2:  # 최대 2문장
                        break
            
            if relevant_sentences:
                return " ".join(relevant_sentences)
            
            # 키워드가 포함된 문장이 없으면 첫 번째 문장 반환
            return sentences[0].strip() if sentences else None
            
        except:
            return None
    
    def _split_into_sentences(self, text):
        """텍스트를 문장으로 분할"""
        # 문장 구분자로 분할
        sentences = re.split(r'[.!?]\s+', text)
        
        # 너무 짧은 문장 제거
        valid_sentences = []
        for sentence in sentences:
            if len(sentence.strip()) > 10:  # 10자 이상인 문장만
                valid_sentences.append(sentence.strip())
        
        return valid_sentences
    
    def _clean_and_format_content(self, content):
        """내용 정리 및 포맷팅 (띄어쓰기 복원 포함)"""
        # 1. 기본 정리
        content = re.sub(r'\s+', ' ', content)
        content = re.sub(r'[^\w\s가-힣,.!?():\-]', '', content)
        content = content.strip()
        
        # 2. 띄어쓰기 복원 (개선된 함수 사용)
        content = self.restore_korean_spacing(content)
        
        return content
    
    def restore_korean_spacing(self, text):
        """한국어 띄어쓰기 복원 함수 - 3단계 접근법"""
        if not text or len(text.strip()) == 0:
            return text
            
        try:
            # 1단계: PyKoSpacing 사용 (가장 좋은 품질)
            if PYKOSPACING_AVAILABLE:
                return spacing(text)
            
            # 2단계: KoNLPy 사용 (형태소 분석 기반)
            elif KONLPY_AVAILABLE:
                return self._restore_spacing_with_konlpy(text)
            
            # 3단계: 패턴 매칭 기반 (라이브러리 없이)
            else:
                return self._restore_spacing_with_patterns(text)
                
        except Exception as e:
            print(f"띄어쓰기 복원 오류: {e}")
            return self._restore_spacing_with_patterns(text)
    
    def _restore_spacing_with_konlpy(self, text):
        """KoNLPy를 사용한 띄어쓰기 복원"""
        try:
            okt = Okt()
            
            # 형태소 분석 후 띄어쓰기 적용 (normalize, stem 파라미터 제거)
            morphs = okt.morphs(text)
            pos_tags = okt.pos(text)
            
            spaced_text = ""
            for i, (morph, pos) in enumerate(pos_tags):
                if i == 0:
                    spaced_text += morph
                else:
                    # 조사, 어미, 접미사는 붙여쓰고 나머지는 띄어쓰기
                    if pos in ['Josa', 'Eomi', 'Suffix']:
                        spaced_text += morph
                    else:
                        spaced_text += " " + morph
            
            return spaced_text
            
        except Exception as e:
            print(f"KoNLPy 처리 중 오류: {e}")
            return self._restore_spacing_with_patterns(text)
    
    def _restore_spacing_with_patterns(self, text):
        """패턴 매칭을 사용한 띄어쓰기 복원 (라이브러리 없이)"""
        if not text:
            return text
        
        # 기본 정리
        text = re.sub(r'\s+', ' ', text).strip()
        
        patterns = [
            # 문장부호 뒤에 띄어쓰기
            (r'([.!?])([가-힣A-Za-z0-9])', r'\1 \2'),
            (r'([,;:])([가-힣A-Za-z0-9])', r'\1 \2'),
            
            # 숫자와 문자 사이
            (r'([0-9])([가-힣])', r'\1 \2'),
            (r'([가-힣])([0-9])', r'\1 \2'),
            
            # 영어와 한글 사이
            (r'([A-Za-z])([가-힣])', r'\1 \2'),
            (r'([가-힣])([A-Za-z])', r'\1 \2'),
            
            # 자주 사용되는 단어들 뒤
            (r'(이다|하다|되다|있다|없다|같다|이며|라고|이라고|한다|된다)([가-힣])', r'\1 \2'),
            (r'(그리고|하지만|그러나|또한|따라서|즉|예를들어|때문에)([가-힣])', r'\1 \2'),
            (r'(경우에|관련하여|대하여|통하여|위하여|의하여)([가-힣])', r'\1 \2'),
            
            # 기술용어와 한글 사이
            (r'(AI|ML|LLM|RAG|API|GPU|CPU|NLP|CNN|RNN|IoT|VR|AR)([가-힣])', r'\1 \2'),
            (r'([가-힣])(AI|ML|LLM|RAG|API|GPU|CPU|NLP|CNN|RNN|IoT|VR|AR)', r'\1 \2'),
            
            # 조사 앞뒤 처리 (조심스럽게)
            (r'([가-힣])(은는이가을를에서의로와과도만큼부터까지마다에게께서)([가-힣A-Z])', r'\1\2 \3'),
            
        ]
        
        result = text
        for pattern, replacement in patterns:
            result = re.sub(pattern, replacement, result)
        
        # 긴 한글 단어 분할 (별도 처리)
        result = self._split_long_korean_words(result)
        
        # 최종 정리
        result = re.sub(r'\s+', ' ', result)
        result = re.sub(r' +([,.!?;:])', r'\1', result)
        
        return result.strip()
    
    def _split_long_korean_words(self, text):
        """긴 한글 단어들을 분할"""
        def split_word(match):
            word = match.group(0)
            if len(word) <= 8:
                return word
                
            # 15자 이상이면 3등분
            if len(word) > 15:
                third = len(word) // 3
                return word[:third] + ' ' + word[third:2*third] + ' ' + word[2*third:]
            # 8자 이상이면 반으로
            else:
                half = len(word) // 2
                return word[:half] + ' ' + word[half:]
        
        # 8자 이상 연속된 한글 단어 찾아서 분할
        return re.sub(r'[가-힣]{8,}', split_word, text)
    
    def _smart_split_word(self, word):
        """긴 단어를 스마트하게 분할 (호환성을 위해 유지)"""
        if len(word) <= 8:
            return word
            
        # 15자 이상이면 3등분
        if len(word) > 15:
            third = len(word) // 3
            return word[:third] + ' ' + word[third:2*third] + ' ' + word[2*third:]
        # 8자 이상이면 반으로
        else:
            half = len(word) // 2
            return word[:half] + ' ' + word[half:]

    def _restore_spacing_advanced(self, text):
        """고급 띄어쓰기 복원 (기존 함수 유지 - 호환성을 위해)"""
        return self.restore_korean_spacing(text)
    
    def _restore_spacing_enhanced(self, text):
        """강화된 띄어쓰기 복원 (라이브러리 없이) - 한국어 특화"""
        if not text or len(text) < 2:
            return text
        
        # 1. 기본 정리
        text = re.sub(r'\s+', ' ', text).strip()
        
        # 2. 한국어 띄어쓰기 패턴들 (실전 특화)
        patterns = [
            # 마침표, 느낌표, 물음표 뒤
            (r'([.!?])([가-힣A-Za-z가-힣])', r'\1 \2'),
            (r'([,;:])([가-힣A-Za-z])', r'\1 \2'),
            
            # 숫자와 한글 사이
            (r'([0-9])([가-힣])', r'\1 \2'),
            (r'([가-힣])([0-9])', r'\1 \2'),
            
            # 영어와 한글 사이  
            (r'([A-Za-z])([가-힣])', r'\1 \2'),
            (r'([가-힣])([A-Za-z])', r'\1 \2'),
            
            # 중요: 자주 등장하는 단어들 뒤에 띄어쓰기
            (r'(이다|하다|되다|있다|없다|같다|이며|라고|이라고)([가-힣])', r'\1 \2'),
            (r'(그리고|하지만|그러나|또한|따라서|즉|예를들어)([가-힣])', r'\1 \2'),
            (r'(때문에|경우에|관련하여|대하여|통하여)([가-힣])', r'\1 \2'),
            
            # 기술용어 분리
            (r'(LLM|RAG|AI|ML|API|GPU|CPU|NLP|CNN|RNN)([가-힣])', r'\1 \2'),
            (r'([가-힣])(LLM|RAG|AI|ML|API|GPU|CPU|NLP|CNN|RNN)', r'\1 \2'),
            
            # 특정 패턴 - 실제 문서에서 자주 보이는 것들
            (r'([가-힣])([CDEFGHIJKLMNOPQRSTUVWXYZ가-힣]{2,})', r'\1 \2'),
            (r'(을수|를수|에대해|에관해|로부터)([가-힣])', r'\1 \2'),
            
            # 조사 앞뒤 (조심스럽게)
            (r'([가-힣])(은는이가을를에서의로와과도만큼부터까지마다)([가-힣A-Z])', r'\1\2 \3'),
        ]
        
        result = text
        for pattern, replacement in patterns:
            result = re.sub(pattern, replacement, result)
        
        # 3. 특별히 긴 단어들 처리 (문서에서 자주 보이는 패턴)
        result = self._break_long_words(result)
        
        # 4. 정리
        result = re.sub(r'\s+', ' ', result)
        result = re.sub(r' +([,.!?;:])', r'\1', result)
        
        return result.strip()
    
    def _break_long_words(self, text):
        """과도하게 긴 연속 단어들을 적절히 분할"""
        # 10자 이상 연속된 한글을 찾아서 중간에 띄어쓰기 추가
        def split_long_korean(match):
            word = match.group(0)
            if len(word) > 15:  # 15자 이상이면 3등분
                third = len(word) // 3
                return word[:third] + ' ' + word[third:2*third] + ' ' + word[2*third:]
            elif len(word) > 8:  # 8자 이상이면 반으로
                half = len(word) // 2
                return word[:half] + ' ' + word[half:]
            return word
        
        # 연속된 한글 패턴 찾기
        text = re.sub(r'[가-힣]{8,}', split_long_korean, text)
        
        return text
    
    def _restore_spacing(self, text):
        """띄어쓰기 복원 함수"""
        if not text:
            return text
        
        # 한국어 띄어쓰기 패턴 적용
        patterns = [
            # 조사 앞에 띄어쓰기
            (r'([가-힣])([은는이가을를에서와과로으로의])([가-힣])', r'\1\2 \3'),
            
            # 마침표, 물음표, 느낌표 뒤에 띄어쓰기  
            (r'([.!?])([가-힣A-Za-z])', r'\1 \2'),
            
            # 쉼표 뒤에 띄어쓰기
            (r'([,])([가-힣A-Za-z])', r'\1 \2'),
            
            # 숫자와 한글 사이
            (r'([0-9])([가-힣])', r'\1 \2'),
            (r'([가-힣])([0-9])', r'\1 \2'),
            
            # 영어와 한글 사이
            (r'([A-Za-z])([가-힣])', r'\1 \2'),
            (r'([가-힣])([A-Za-z])', r'\1 \2'),
            
            # 특정 단어들 뒤에 띄어쓰기
            (r'(이다|있다|없다|하다|되다|같다)([가-힣])', r'\1 \2'),
            (r'(그리고|하지만|그러나|따라서|또한)([가-힣])', r'\1 \2'),
            
            # 자주 사용되는 접속어들
            (r'(LLM|RAG|AI|ML)([가-힣])', r'\1 \2'),
            (r'([가-힣])(LLM|RAG|AI|ML)', r'\1 \2'),
        ]
        
        result = text
        for pattern, replacement in patterns:
            result = re.sub(pattern, replacement, result)
        
        # 과도한 띄어쓰기 정리
        result = re.sub(r'\s+', ' ', result)
        
    def _restore_spacing_advanced(self, text):
        """고급 띄어쓰기 복원 (라이브러리 사용)"""
        try:
            # python-spacing 라이브러리 사용 (설치된 경우)
            from spacing import spacing
            return spacing(text)
        except ImportError:
            try:
                # khaiii 사용 (설치된 경우)
                from khaiii import KhaiiiApi
                api = KhaiiiApi()
                spaced_text = ""
                
                for word in api.analyze(text):
                    spaced_text += word.lex + " "
                
                return spaced_text.strip()
            except ImportError:
                # 라이브러리가 없으면 기본 방법 사용
                return self._restore_spacing(text)
        except Exception:
            # 오류 발생시 기본 방법 사용
            return self._restore_spacing(text)
    
    def _smart_truncate(self, text, max_length):
        """스마트한 텍스트 자르기 (문장 단위로)"""
        if len(text) <= max_length:
            return text
        
        # max_length 근처에서 문장 끝을 찾아 자르기
        truncated = text[:max_length]
        
        # 마지막 완전한 문장 찾기
        last_period = truncated.rfind('.')
        last_exclamation = truncated.rfind('!')
        last_question = truncated.rfind('?')
        
        last_sentence_end = max(last_period, last_exclamation, last_question)
        
        if last_sentence_end > max_length * 0.7:  # 70% 이상이면 문장 단위로 자르기
            return truncated[:last_sentence_end + 1]
        else:
            # 문장 끝이 너무 앞에 있으면 단어 단위로 자르기
            last_space = truncated.rfind(' ')
            if last_space > 0:
                return truncated[:last_space] + "..."
            else:
                return truncated + "..."
    
    def get_document_count(self):
        """저장된 문서 청크 개수"""
        return len(self.documents)
    
    def get_document_stats(self):
        """문서 통계 정보"""
        if not self.documents:
            return {"총 문서": 0, "평균 길이": 0}
        
        sources = {}
        total_length = 0
        
        for doc in self.documents:
            source = doc.metadata.get('source', '알 수 없음')
            sources[source] = sources.get(source, 0) + 1
            total_length += len(doc.page_content)
        
        return {
            "총 청크": len(self.documents),
            "문서 종류": len(sources),
            "평균 청크 길이": total_length // len(self.documents),
            "문서별 청크": sources
        }

# 전역 RAG 엔진 인스턴스
rag_engine = SimpleRAGEngine()